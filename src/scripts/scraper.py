import sys
import json
import requests
import os
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import datetime

# Define color styles for span classes
color_styles = {
    "c1": "color: red;",
    "c2": "color: green;",
    "c3": "color: blue;",
    "c4": "color: purple;",
    "c5": "color: orange;"
}

def generate_urls(base_url, book_number, start_page, end_page):
    """Generates a list of URLs based on the book number and page range."""
    urls = []
    for page in range(start_page-1, end_page):
        url = base_url.format(bookNumber=book_number, pageNumber=page)
        urls.append(url)
    return urls

def scrape_page_content(url):
    """Scrapes data and HTML content from a given URL."""
    try:
        response = requests.get(url)
        if response.status_code != 200:
            return {"error": f"Failed to fetch the page. Status code: {response.status_code}"}

        soup = BeautifulSoup(response.text, 'html.parser')

        # Extract fields such as Hadith No., Volume, Book Page, Chapter, and Question
        hadith_no = soup.find("input", id="fld_specialNum_top").get("value", "Not found") if soup.find("input", id="fld_specialNum_top") else "Hadith No. not found."
        volume_no = soup.find("button", {"data-toggle": "dropdown"}).get_text(strip=True) if soup.find("button", {"data-toggle": "dropdown"}) else "Volume not found."
        book_page = soup.find("input", id="fld_goto_top").get("value", "Not found") if soup.find("input", id="fld_goto_top") else "Book Page not found."
        
        # Extract Chapter and Question
        main_div_info = soup.find("div", class_="size-12")
        chapter, question = "Chapter text not found.", "Question text not found."
        if main_div_info:
            text_black_spans = main_div_info.find_all("span", class_="text-black")
            if len(text_black_spans) > 1:
                chapter = text_black_spans[1].get_text(strip=True)
            if len(text_black_spans) > 2:
                question = text_black_spans[2].get_text(strip=True)

        # Extract Issue, Fatwa, and Hamesh
        main_div = soup.find("div", class_="nass margin-top-10")
        issue_html, fatwa_html, hamesh = "", "", ""
        if main_div:
            paragraphs = main_div.find_all('p')
            
            # Format Issue with color styling
            if len(paragraphs) >= 2:
                issue_paragraphs = []
                for para in paragraphs[:2]:
                    p_soup = BeautifulSoup(str(para), 'html.parser')
                    for span in p_soup.find_all("span"):
                        span_class = span.get("class", [None])[0]
                        if span_class in color_styles:
                            span['style'] = color_styles[span_class]
                    issue_paragraphs.append(str(p_soup))
                issue_html = "".join(issue_paragraphs)

            # Format Fatwa with color styling
            fatwa_texts = []
            for i in range(2, len(paragraphs)):
                if 'hamesh' in paragraphs[i].get('class', []):
                    hamesh = paragraphs[i].get_text(separator="\n", strip=True).replace("\n", "<br>")
                    break
                p_content = str(paragraphs[i])
                p_soup = BeautifulSoup(p_content, 'html.parser')
                for span in p_soup.find_all("span"):
                    span_class = span.get("class", [None])[0]
                    if span_class in color_styles:
                        span['style'] = color_styles[span_class]
                fatwa_texts.append(str(p_soup))
            fatwa_html = "".join(fatwa_texts)
        
        # Prepare HTML content with table formatting
        html_content_with_table = f"""
        <html lang="ar" dir="rtl">
        <head>
        <meta charset="UTF-8">
        <style>
            body {{ font-family: Arial, sans-serif; direction: rtl; }}
            p {{ line-height: 1.6; }}
            table {{ width: 100%; border-collapse: collapse; }}
            th, td {{ border: 1px solid #ddd; padding: 8px; text-align: right; }}
            th {{ background-color: #f2f2f2; }}
        </style>
        </head>
        <body>
        <table>
            <tr>
                <th>Hadith No.</th>
                <td>{hadith_no}</td>
            </tr>
            <tr>
                <th>Volume</th>
                <td>{volume_no}</td>
            </tr>
            <tr>
                <th>Book Page</th>
                <td>{book_page}</td>
            </tr>
            <tr>
                <th>Chapter</th>
                <td>{chapter}</td>
            </tr>
            <tr>
                <th>Question</th>
                <td>{question}</td>
            </tr>
            <tr>
                <th>Issue</th>
                <td>{issue_html}</td>
            </tr>
            <tr>
                <th>Fatwa</th>
                <td>{fatwa_html}</td>
            </tr>
            <tr>
                <th>Hamesh</th>
                <td>{hamesh}</td>
            </tr>
        </table>
        </body>
        </html>
        """
        
        # Prepare HTML content without table (only content)
        html_content_without_table = f"""
        <html lang="ar" dir="rtl">
        <head>
        <meta charset="UTF-8">
        <style>
            body {{ font-family: Arial, sans-serif; direction: rtl; }}
            p {{ line-height: 1.6; }}
        </style>
        </head>
        <body>
        {issue_html}
        {fatwa_html}
        {hamesh}
        <div style="page-break-after: always;">..............</div>
        </body>
        </html>
        """
        
        return {
            "url": url,
            "hadith_no": hadith_no,
            "volume_no": volume_no,
            "book_page": book_page,
            "chapter": chapter,
            "question": question,
            "issue": issue_html,
            "fatwa": fatwa_html,
            "hamesh": hamesh,
            "html_content_with_table": html_content_with_table,
            "html_content_without_table": html_content_without_table
        }

    except Exception as e:
        return {"error": f"Error: {str(e)}"}

def save_html_file(content, folder_name, file_name):
    """Saves HTML content to a file."""
    os.makedirs(folder_name, exist_ok=True)
    file_path = os.path.join(folder_name, file_name)
    with open(file_path, "w", encoding="utf-8") as file:
        file.write(content)
    print(f"HTML file saved at '{file_path}'", file=sys.stderr)

def combine_html_with_table(input_dir, output_dir, run_id):
    """Combines all HTML files with table data into a single document."""
    os.makedirs(output_dir, exist_ok=True)
    combined_path = os.path.join(output_dir, f"combined_with_table_{run_id}.html")
    
    with open(combined_path, "w", encoding="utf-8") as outfile:
        outfile.write('<html lang="ar" dir="rtl">\n<head><meta charset="UTF-8"></head>\n<body>\n')
        
        for filename in sorted(os.listdir(input_dir)):
            if filename.endswith('.html') and "_with_table" in filename:
                filepath = os.path.join(input_dir, filename)
                with open(filepath, "r", encoding="utf-8") as infile:
                    content = infile.read()
                    outfile.write(content)
        
        outfile.write('</body>\n</html>')
    return combined_path

def combine_html_without_table(input_dir, output_dir, run_id):
    """Combines all HTML files without table data into a single document."""
    os.makedirs(output_dir, exist_ok=True)
    combined_path = os.path.join(output_dir, f"combined_without_table_{run_id}.html")
    
    with open(combined_path, "w", encoding="utf-8") as outfile:
        outfile.write('<html lang="ar" dir="rtl">\n<head><meta charset="UTF-8"></head>\n<body>\n')
        
        for filename in sorted(os.listdir(input_dir)):
            if filename.endswith('.html') and "_without_table" in filename:
                filepath = os.path.join(input_dir, filename)
                with open(filepath, "r", encoding="utf-8") as infile:
                    content = infile.read()
                    outfile.write(content)
                    outfile.write('<div style="page-break-after: always;">..............</div>')  # Add page break
        
        outfile.write('</body>\n</html>')
    return combined_path

def add_page_break(doc):
    """Adds a page break to the document."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    break_element = OxmlElement("w:br")
    break_element.set(qn("w:type"), "page")
    run._r.append(break_element)

def html_to_docx(html_path, docx_path):
    """Converts combined HTML file (without table data) to a formatted DOCX document."""
    doc = Document()
    
    with open(html_path, "r", encoding="utf-8") as f:
        soup = BeautifulSoup(f, "html.parser")
        
        # Process all elements in the body
        body = soup.find("body")
        if body:
            for element in body.find_all(["h3", "p", "small"]):
                if element.name == "h3":
                    heading = doc.add_heading(element.text, level=3)
                    heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                elif element.name == "p":
                    para = doc.add_paragraph()
                    run = para.add_run(element.text)
                    para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                elif element.name == "small":
                    para = doc.add_paragraph(element.text)
                    para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    para.runs[0].font.size = Pt(10)
            
            # Add a page break after each page's content
            add_page_break(doc)
    
    doc.save(docx_path)

def html_to_doc(html_path, doc_path):
    """Converts combined HTML file (without table data) to a DOC file."""
    temp_docx_path = doc_path.replace(".doc", ".docx")
    html_to_docx(html_path, temp_docx_path)
    os.rename(temp_docx_path, doc_path)

def main(base_url, book_number, start_page, end_page):
    """Main function to process URLs and generate outputs."""
    # Generate unique Run ID using current timestamp
    run_id = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    
    # Create output directories
    scraped_dir = f"{book_number}_{start_page}-{end_page}"
    combined_dir = "combined_output"
    
    # Generate URLs dynamically
    urls = generate_urls(base_url, book_number, start_page, end_page)
    
    # Scrape each URL
    results = []
    for url in urls:
        data = scrape_page_content(url)
        if "error" in data:
            results.append(data)
            continue
        
        # Generate unique filenames with Run ID
        page_suffix = url.split('/')[-1].replace('#', '_')
        filename_with_table = f"page_{page_suffix}_{run_id}_with_table.html"
        filename_without_table = f"page_{page_suffix}_{run_id}_without_table.html"
        
        # Save HTML files
        save_html_file(data["html_content_with_table"], scraped_dir, filename_with_table)
        save_html_file(data["html_content_without_table"], scraped_dir, filename_without_table)
        results.append(data)
    
    # Combine HTML files with table data
    combined_html_with_table = combine_html_with_table(scraped_dir, combined_dir, run_id)
    
    # Combine HTML files without table data
    combined_html_without_table = combine_html_without_table(scraped_dir, combined_dir, run_id)
    
    # Convert the new combined HTML (without table data) to DOCX
    output_docx = os.path.join(combined_dir, f"scraped output_{run_id}.docx")
    html_to_docx(combined_html_without_table, output_docx)
    
    # Convert the new combined HTML (without table data) to DOC
    output_doc = os.path.join(combined_dir, f"scraped output_{run_id}.doc")
    html_to_doc(combined_html_without_table, output_doc)
    
    # Output JSON result
    result = {
        "status": "success",
        "output_docx": output_docx,
        "output_doc": output_doc,
        "scraped_data": results
    }
    print(json.dumps(result))

if __name__ == "__main__":
    if len(sys.argv) != 5:
        print(json.dumps({"error": "Usage: python script.py <base_url> <book_number> <start_page> <end_page>"}))
        sys.exit(1)
    
    base_url = sys.argv[1]
    book_number = int(sys.argv[2])
    start_page = int(sys.argv[3])
    end_page = int(sys.argv[4])
    
    main(base_url, book_number, start_page, end_page)